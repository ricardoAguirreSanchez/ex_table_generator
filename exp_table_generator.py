"""
Exp Table Generator — versión GUI portable.

Lee un Word template para detectar columnas, anchos y formato.
Mapea automáticamente las columnas del template a las del Excel.
Permite ajustar el mapeo manualmente antes de generar.
"""

import sys
import re
import unicodedata
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
from pathlib import Path

from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

MONTH_MAP = {
    "Enero": "ene", "Febrero": "feb", "Marzo": "mar", "Abril": "abr",
    "Mayo": "may", "Junio": "jun", "Julio": "jul", "Agosto": "ago",
    "Septiembre": "sep", "Octubre": "oct", "Noviembre": "nov", "Diciembre": "dic",
}

COUNTRY_KEYWORDS = {
    "Argentina": "Argentina", "argentina": "Argentina",
    "Perú": "Perú", "Peru": "Perú",
    "Colombia": "Colombia", "Chile": "Chile", "Bolivia": "Bolivia",
    "Ecuador": "Ecuador", "Brasil": "Brasil", "Paraguay": "Paraguay",
    "Uruguay": "Uruguay", "México": "México", "Mexico": "México",
    "Panamá": "Panamá", "Panama": "Panamá",
    "Costa Rica": "Costa Rica", "Honduras": "Honduras",
    "El Salvador": "El Salvador", "Guatemala": "Guatemala",
    "Nicaragua": "Nicaragua", "Venezuela": "Venezuela",
    "República Dominicana": "República Dominicana",
}

ENTITY_COUNTRY_HINTS = {
    "Nación": "Argentina", "Nacion": "Argentina",
    "Buenos Aires": "Argentina", "CABA": "Argentina",
    "Provincia de": "Argentina",
    "Ministerio de Seguridad": "Argentina",
    "Banco Hipotecario": "Argentina",
}

ALIGN_MAP = {
    "CENTER (1)": WD_ALIGN_PARAGRAPH.CENTER,
    "JUSTIFY (3)": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "LEFT (0)": WD_ALIGN_PARAGRAPH.LEFT,
    "RIGHT (2)": WD_ALIGN_PARAGRAPH.RIGHT,
}

SPECIAL_SOURCES = ["(auto-incremento)", "(extraer país)"]


def base_dir() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent


# ---------------------------------------------------------------------------
# Text helpers
# ---------------------------------------------------------------------------

def normalize(text: str) -> str:
    """Remove accents, lowercase, strip non-alpha."""
    text = unicodedata.normalize("NFD", text.lower())
    text = "".join(c for c in text if unicodedata.category(c) != "Mn")
    return re.sub(r"[^a-z0-9 ]", " ", text).strip()


def word_overlap(a: str, b: str) -> float:
    """Score how much two header strings overlap (0-1)."""
    wa = set(normalize(a).split())
    wb = set(normalize(b).split())
    if not wa or not wb:
        return 0.0
    return len(wa & wb) / max(len(wa), len(wb))


def convert_date(date_str: str) -> str:
    if not date_str:
        return ""
    parts = str(date_str).strip().split()
    if len(parts) == 2:
        abbr = MONTH_MAP.get(parts[0], parts[0][:3].lower())
        return f"{abbr}-{parts[1][-2:]}"
    return str(date_str)


def extract_country(text: str) -> str:
    if not text:
        return ""
    t = str(text)
    for kw, country in COUNTRY_KEYWORDS.items():
        if kw in t:
            return country
    for kw, country in ENTITY_COUNTRY_HINTS.items():
        if kw in t:
            return country
    return ""


def col_letter_to_index(letter: str) -> int:
    return ord(letter.upper()) - 64


# ---------------------------------------------------------------------------
# Template reader
# ---------------------------------------------------------------------------

def read_template(template_path: str) -> dict:
    """Extract column metadata from a Word template's first table."""
    doc = Document(template_path)
    if not doc.tables:
        raise ValueError("El template Word no contiene tablas.")

    table = doc.tables[0]

    # Title from paragraph before the table
    title = ""
    for p in doc.paragraphs:
        if p.text.strip():
            title = p.text.strip()
            break

    # Grid widths
    tblGrid = table._tbl.find(qn("w:tblGrid"))
    grid_widths = []
    if tblGrid is not None:
        for gc in tblGrid.findall(qn("w:gridCol")):
            grid_widths.append(int(gc.get(qn("w:w"))))

    # Page setup
    section = doc.sections[0]
    page_info = {
        "width": section.page_width,
        "height": section.page_height,
        "orientation": int(section.orientation),
        "left_margin": section.left_margin,
        "right_margin": section.right_margin,
        "top_margin": section.top_margin,
        "bottom_margin": section.bottom_margin,
    }

    columns = []
    has_data_row = len(table.rows) > 1

    for i, header_cell in enumerate(table.rows[0].cells):
        header = header_cell.text.strip()
        width = grid_widths[i] if i < len(grid_widths) else 1500

        data_bold = False
        data_align = "CENTER (1)"
        if has_data_row:
            data_cell = table.rows[1].cells[i]
            for p in data_cell.paragraphs:
                data_align = str(p.alignment) if p.alignment else "CENTER (1)"
                for r in p.runs:
                    if r.bold:
                        data_bold = True
                break

        columns.append({
            "header": header,
            "width": width,
            "bold": data_bold,
            "align": data_align,
        })

    return {
        "title": title,
        "page": page_info,
        "columns": columns,
    }


# ---------------------------------------------------------------------------
# Excel reader
# ---------------------------------------------------------------------------

def read_excel_headers(excel_path: str, sheet: str = "ESP", header_row: int = 3) -> dict:
    """Return {col_letter: header_text} for non-empty columns."""
    wb = load_workbook(excel_path, data_only=True, read_only=True)
    if sheet not in wb.sheetnames:
        raise ValueError(f"Hoja '{sheet}' no encontrada. Disponibles: {wb.sheetnames}")
    ws = wb[sheet]

    headers = {}
    for col in range(1, 27):
        v1 = None
        v2 = None
        for row_idx, row in enumerate(ws.iter_rows(min_row=header_row, max_row=header_row + 1,
                                                     min_col=col, max_col=col), start=0):
            for cell in row:
                if row_idx == 0:
                    v1 = cell.value
                else:
                    v2 = cell.value

        combined = ""
        if v1 and v2:
            combined = f"{v1} / {v2}"
        elif v1:
            combined = str(v1)
        elif v2:
            combined = str(v2)

        if combined.strip():
            letter = chr(64 + col)
            headers[letter] = combined.strip()

    wb.close()
    return headers


def peek_excel_rows(excel_path: str, sheet: str = "ESP", header_row: int = 3) -> str:
    wb = load_workbook(excel_path, data_only=True, read_only=True)
    if sheet not in wb.sheetnames:
        wb.close()
        return f"ERROR: Hoja '{sheet}' no encontrada."
    ws = wb[sheet]

    lines = [f"Hoja: {sheet}\n"]
    shown = 0
    for row in ws.iter_rows(min_row=header_row + 2, max_col=2, values_only=False):
        for cell in row:
            if cell.value:
                lines.append(f"  Fila {cell.row}: {str(cell.value)[:90]}")
                shown += 1
        if shown >= 20:
            lines.append("  ...")
            break

    wb.close()
    return "\n".join(lines)


def read_excel_data(excel_path: str, row_numbers: list[int], mapping: list[dict],
                    sheet: str = "ESP") -> list[dict]:
    """Read specific rows using the column mapping."""
    wb = load_workbook(excel_path, data_only=True)
    try:
        ws = wb[sheet]

        rows = []
        for seq, row_num in enumerate(row_numbers, start=1):
            if row_num < 1 or row_num > ws.max_row:
                continue
            row_data = {}
            for m in mapping:
                header = m["header"]
                source = m["source"]

                if source == "(auto-incremento)":
                    row_data[header] = str(seq)
                elif source == "(extraer país)":
                    from_col = m.get("from_col", "D")
                    raw = str(ws.cell(row=row_num, column=col_letter_to_index(from_col)).value or "")
                    row_data[header] = extract_country(raw)
                else:
                    col_idx = col_letter_to_index(source)
                    raw = ws.cell(row=row_num, column=col_idx).value
                    val = str(raw) if raw is not None else ""
                    if m.get("format") == "short_date":
                        val = convert_date(val)
                    row_data[header] = val

            rows.append(row_data)
        return rows
    finally:
        wb.close()


# ---------------------------------------------------------------------------
# Auto-mapping heuristics
# ---------------------------------------------------------------------------

DATE_KEYWORDS_START = {"inicio", "desde", "from", "start"}
DATE_KEYWORDS_END = {"fin", "hasta", "end", "until"}

def auto_map(template_cols: list[dict], excel_headers: dict) -> list[dict]:
    """Guess the best Excel column for each template column."""
    mapping = []
    used = set()

    for col in template_cols:
        header_norm = normalize(col["header"])
        best_source = ""
        best_score = 0.0
        fmt = ""

        # Special: "No." / "#" -> auto-increment
        if header_norm in ("no", "no.", "n", "#", "numero"):
            mapping.append({**col, "source": "(auto-incremento)", "format": ""})
            continue

        # Special: "País" / "Country"
        if header_norm in ("pais", "country", "paises"):
            # Find the entity column to extract from
            entity_col = "D"
            for letter, eh in excel_headers.items():
                if "entidad" in normalize(eh) or "contratante" in normalize(eh):
                    entity_col = letter
                    break
            mapping.append({**col, "source": "(extraer país)", "from_col": entity_col, "format": ""})
            continue

        # Date detection
        is_start_date = bool(DATE_KEYWORDS_START & set(header_norm.split()))
        is_end_date = bool(DATE_KEYWORDS_END & set(header_norm.split()))

        for letter, excel_header in excel_headers.items():
            if letter in used:
                continue
            score = word_overlap(col["header"], excel_header)

            # Boost date matching
            eh_norm = normalize(excel_header)
            if is_start_date and ("desde" in eh_norm or "inicio" in eh_norm or "from" in eh_norm):
                score += 0.4
            if is_end_date and ("hasta" in eh_norm or "fin" in eh_norm or "until" in eh_norm):
                score += 0.4

            if score > best_score:
                best_score = score
                best_source = letter

        if is_start_date or is_end_date:
            fmt = "short_date"

        if best_source and best_score > 0.1:
            used.add(best_source)
            mapping.append({**col, "source": best_source, "format": fmt})
        else:
            mapping.append({**col, "source": "", "format": fmt})

    return mapping


# ---------------------------------------------------------------------------
# Word generation
# ---------------------------------------------------------------------------

def build_document(data_rows: list[dict], template_info: dict, mapping: list[dict],
                   output_path: str):
    doc = Document()
    page = template_info["page"]
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE if page["orientation"] == 1 else WD_ORIENT.PORTRAIT
    section.page_width = page["width"]
    section.page_height = page["height"]
    section.left_margin = page["left_margin"]
    section.right_margin = page["right_margin"]
    section.top_margin = page["top_margin"]
    section.bottom_margin = page["bottom_margin"]

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(template_info.get("title", ""))
    run.bold = True
    run.font.size = Pt(11)

    num_cols = len(mapping)
    widths = [m.get("width", 1500) for m in mapping]

    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)

    # Borders
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        table._tbl.insert(0, tblPr)
    tblPr.append(parse_xml(f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>
    </w:tblBorders>'''))

    # Grid
    tblGrid = table._tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = parse_xml(f'<w:tblGrid {nsdecls("w")}></w:tblGrid>')
        table._tbl.insert(1, tblGrid)
    else:
        for child in list(tblGrid):
            tblGrid.remove(child)
    for w in widths:
        tblGrid.append(parse_xml(f'<w:gridCol {nsdecls("w")} w:w="{w}"/>'))

    # Header row
    for i, m in enumerate(mapping):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = p._p.get_or_add_pPr()
        pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:after="0" w:line="240" w:lineRule="auto"/>'))
        r = p.add_run(m["header"])
        r.bold = True
        r.font.size = Pt(11)
        tcPr = cell._tc.get_or_add_tcPr()
        tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="BFBFBF" w:val="clear"/>'))
        tcPr.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>'))
        tcPr.append(parse_xml(f'<w:tcW {nsdecls("w")} w:w="{widths[i]}" w:type="dxa"/>'))

    # Data rows
    for row_idx, row_data in enumerate(data_rows):
        for col_idx, m in enumerate(mapping):
            value = row_data.get(m["header"], "")
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            pPr = p._p.get_or_add_pPr()
            pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:after="0" w:line="240" w:lineRule="auto"/>'))
            align_key = m.get("align", "CENTER (1)")
            p.alignment = ALIGN_MAP.get(align_key, WD_ALIGN_PARAGRAPH.CENTER)
            tcPr = cell._tc.get_or_add_tcPr()
            tcPr.append(parse_xml(f'<w:tcW {nsdecls("w")} w:w="{widths[col_idx]}" w:type="dxa"/>'))
            r = p.add_run(value)
            r.font.size = Pt(11)
            if m.get("bold"):
                r.bold = True

    doc.save(output_path)


# ---------------------------------------------------------------------------
# GUI
# ---------------------------------------------------------------------------

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Exp Table Generator")
        self.geometry("850x700")
        self.resizable(True, True)

        self.template_path = tk.StringVar()
        self.excel_path = tk.StringVar()
        self.sheet_var = tk.StringVar(value="ESP")
        self.header_row_var = tk.StringVar(value="3")

        self.template_info = None
        self.excel_headers = {}
        self.mapping_widgets = []

        self._build_ui()

    # ----- UI build -----

    def _build_ui(self):
        pad = {"padx": 8, "pady": 4}

        # --- Step 1: Template ---
        f1 = tk.LabelFrame(self, text="1. Word Template (modelo)", **pad)
        f1.pack(fill="x", **pad)
        tk.Entry(f1, textvariable=self.template_path, width=70).pack(side="left", padx=5, pady=5, fill="x", expand=True)
        tk.Button(f1, text="Buscar...", command=self._pick_template).pack(side="right", padx=5, pady=5)

        # --- Step 2: Excel ---
        f2 = tk.LabelFrame(self, text="2. Excel (datos fuente)", **pad)
        f2.pack(fill="x", **pad)

        f2_top = tk.Frame(f2)
        f2_top.pack(fill="x")
        tk.Entry(f2_top, textvariable=self.excel_path, width=55).pack(side="left", padx=5, pady=5, fill="x", expand=True)
        tk.Button(f2_top, text="Buscar...", command=self._pick_excel).pack(side="right", padx=5, pady=5)

        f2_opts = tk.Frame(f2)
        f2_opts.pack(fill="x", padx=5)
        tk.Label(f2_opts, text="Hoja:").pack(side="left")
        tk.Entry(f2_opts, textvariable=self.sheet_var, width=12).pack(side="left", padx=3)
        tk.Label(f2_opts, text="Fila encabezado:").pack(side="left", padx=(15, 0))
        tk.Entry(f2_opts, textvariable=self.header_row_var, width=5).pack(side="left", padx=3)
        tk.Button(f2_opts, text="Cargar", command=self._load_excel).pack(side="left", padx=10)

        # --- Preview ---
        f_preview = tk.LabelFrame(self, text="Vista previa del Excel", **pad)
        f_preview.pack(fill="both", expand=True, **pad)
        self.preview_text = scrolledtext.ScrolledText(f_preview, height=6, state="disabled", font=("Consolas", 9))
        self.preview_text.pack(fill="both", expand=True, padx=5, pady=5)

        # --- Step 3: Mapping ---
        f3 = tk.LabelFrame(self, text="3. Mapeo de columnas (Template → Excel)", **pad)
        f3.pack(fill="x", **pad)
        self.mapping_frame = tk.Frame(f3)
        self.mapping_frame.pack(fill="x", padx=5, pady=5)
        self.mapping_hint = tk.Label(f3, text="Cargá un template y un Excel para ver el mapeo.",
                                     fg="gray", anchor="w")
        self.mapping_hint.pack(fill="x", padx=5)

        # --- Step 4: Rows ---
        f4 = tk.LabelFrame(self, text="4. Filas a incluir (separadas por coma, espacio, o rango con guión)", **pad)
        f4.pack(fill="x", **pad)
        self.rows_entry = tk.Entry(f4, width=60)
        self.rows_entry.insert(0, "50, 51")
        self.rows_entry.pack(side="left", padx=5, pady=5, fill="x", expand=True)

        # --- Generate ---
        f_gen = tk.Frame(self)
        f_gen.pack(fill="x", **pad)
        tk.Button(f_gen, text="Generar Word", command=self._generate,
                  bg="#4CAF50", fg="white", font=("Arial", 12, "bold"),
                  height=2, width=22).pack(pady=8)

        self.status_var = tk.StringVar(value="Comenzá seleccionando un template Word y un Excel.")
        tk.Label(self, textvariable=self.status_var, anchor="w", fg="#333").pack(fill="x", padx=10, pady=(0, 5))

    # ----- File pickers -----

    def _pick_template(self):
        path = filedialog.askopenfilename(
            title="Seleccionar template Word",
            filetypes=[("Word", "*.docx"), ("Todos", "*.*")],
            initialdir=str(base_dir() / "Modelo"),
        )
        if path:
            self.template_path.set(path)
            try:
                self.template_info = read_template(path)
                cols = [c["header"] for c in self.template_info["columns"]]
                self.status_var.set(f"Template cargado: {len(cols)} columnas detectadas.")
                self._try_build_mapping()
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo leer el template:\n{e}")

    def _pick_excel(self):
        path = filedialog.askopenfilename(
            title="Seleccionar archivo Excel",
            filetypes=[("Excel", "*.xlsx *.xls"), ("Todos", "*.*")],
            initialdir=str(base_dir() / "Modelo"),
        )
        if path:
            self.excel_path.set(path)
            self._load_excel()

    def _load_excel(self):
        path = self.excel_path.get().strip()
        if not path:
            return
        sheet = self.sheet_var.get().strip()
        header_row = int(self.header_row_var.get().strip() or "3")
        try:
            self.excel_headers = read_excel_headers(path, sheet, header_row)
            preview = peek_excel_rows(path, sheet, header_row)
            self.preview_text.config(state="normal")
            self.preview_text.delete("1.0", "end")
            self.preview_text.insert("1.0", preview)
            self.preview_text.config(state="disabled")
            self.status_var.set(f"Excel cargado: {len(self.excel_headers)} columnas detectadas.")
            self._try_build_mapping()
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo leer el Excel:\n{e}")

    # ----- Mapping UI -----

    def _try_build_mapping(self):
        if not self.template_info or not self.excel_headers:
            return

        mapping = auto_map(self.template_info["columns"], self.excel_headers)

        for w in self.mapping_frame.winfo_children():
            w.destroy()
        self.mapping_widgets.clear()

        # Header labels
        tk.Label(self.mapping_frame, text="Columna del template", font=("Arial", 9, "bold"),
                 anchor="w", width=35).grid(row=0, column=0, padx=3, sticky="w")
        tk.Label(self.mapping_frame, text="→", font=("Arial", 9, "bold")).grid(row=0, column=1)
        tk.Label(self.mapping_frame, text="Columna del Excel", font=("Arial", 9, "bold"),
                 anchor="w", width=35).grid(row=0, column=2, padx=3, sticky="w")
        tk.Label(self.mapping_frame, text="Formato", font=("Arial", 9, "bold"),
                 anchor="w", width=12).grid(row=0, column=3, padx=3, sticky="w")

        # Options for dropdowns
        excel_options = SPECIAL_SOURCES.copy()
        for letter, header in sorted(self.excel_headers.items()):
            display = f"{letter}: {header[:40]}"
            excel_options.append(display)

        format_options = ["(ninguno)", "short_date"]

        for i, m in enumerate(mapping):
            row = i + 1
            tk.Label(self.mapping_frame, text=m["header"], anchor="w",
                     width=35).grid(row=row, column=0, padx=3, sticky="w")
            tk.Label(self.mapping_frame, text="→").grid(row=row, column=1)

            # Find current selection
            source = m.get("source", "")
            current_display = ""
            if source in SPECIAL_SOURCES:
                current_display = source
            elif source and source in self.excel_headers:
                current_display = f"{source}: {self.excel_headers[source][:40]}"

            combo = ttk.Combobox(self.mapping_frame, values=excel_options, width=38, state="readonly")
            if current_display:
                combo.set(current_display)
            combo.grid(row=row, column=2, padx=3, pady=2, sticky="w")

            fmt_combo = ttk.Combobox(self.mapping_frame, values=format_options, width=12, state="readonly")
            fmt_val = m.get("format", "")
            fmt_combo.set(fmt_val if fmt_val else "(ninguno)")
            fmt_combo.grid(row=row, column=3, padx=3, pady=2, sticky="w")

            self.mapping_widgets.append({
                "header": m["header"],
                "combo": combo,
                "fmt_combo": fmt_combo,
                "width": m.get("width", 1500),
                "bold": m.get("bold", False),
                "align": m.get("align", "CENTER (1)"),
            })

        self.mapping_hint.config(text=f"Se auto-mapearon {len(mapping)} columnas. Ajustá si es necesario.")

    def _get_final_mapping(self) -> list[dict]:
        mapping = []
        for w in self.mapping_widgets:
            raw = w["combo"].get().strip()
            fmt_raw = w["fmt_combo"].get().strip()
            fmt = "" if fmt_raw == "(ninguno)" else fmt_raw

            if not raw:
                continue

            m = {
                "header": w["header"],
                "width": w["width"],
                "bold": w["bold"],
                "align": w["align"],
                "format": fmt,
            }

            if raw == "(auto-incremento)":
                m["source"] = "(auto-incremento)"
            elif raw == "(extraer país)":
                m["source"] = "(extraer país)"
                # Find entity column for extraction
                for letter, eh in self.excel_headers.items():
                    if "entidad" in normalize(eh) or "contratante" in normalize(eh):
                        m["from_col"] = letter
                        break
                else:
                    m["from_col"] = "D"
            else:
                col_letter = raw.split(":")[0].strip()
                m["source"] = col_letter

            mapping.append(m)
        return mapping

    # ----- Row parsing -----

    def _parse_rows(self) -> list[int]:
        raw = self.rows_entry.get().strip().replace(",", " ").replace(";", " ")
        rows = []
        for p in raw.split():
            p = p.strip()
            if not p:
                continue
            if "-" in p and not p.startswith("-"):
                a, b = p.split("-", 1)
                rows.extend(range(int(a), int(b) + 1))
            else:
                rows.append(int(p))
        return rows

    # ----- Generate -----

    def _generate(self):
        if not self.template_info:
            messagebox.showwarning("Atención", "Seleccioná un template Word primero.")
            return
        excel = self.excel_path.get().strip()
        if not excel or not Path(excel).exists():
            messagebox.showwarning("Atención", "Seleccioná un archivo Excel válido.")
            return
        if not self.mapping_widgets:
            messagebox.showwarning("Atención", "No hay mapeo de columnas. Cargá template y Excel.")
            return

        try:
            row_numbers = self._parse_rows()
        except ValueError:
            messagebox.showerror("Error", "Números de fila inválidos.\nEjemplo: 50, 51  o  10-15")
            return
        if not row_numbers:
            messagebox.showwarning("Atención", "Ingresá al menos un número de fila.")
            return

        mapping = self._get_final_mapping()
        if not mapping:
            messagebox.showwarning("Atención", "Todas las columnas están sin mapear.")
            return

        output = filedialog.asksaveasfilename(
            title="Guardar Word como...",
            defaultextension=".docx",
            filetypes=[("Word", "*.docx")],
            initialfile=f"Tabla_filas_{'_'.join(str(r) for r in row_numbers)}.docx",
            initialdir=str(Path(excel).parent),
        )
        if not output:
            return

        try:
            self.status_var.set("Generando...")
            self.update_idletasks()

            sheet = self.sheet_var.get().strip()
            data = read_excel_data(excel, row_numbers, mapping, sheet)
            if not data:
                messagebox.showwarning("Atención", "No se obtuvieron datos para esas filas.")
                return

            build_document(data, self.template_info, mapping, output)
            self.status_var.set(f"Listo: {output}")
            messagebox.showinfo("Éxito", f"Documento generado con {len(data)} filas:\n\n{output}")
        except Exception as e:
            self.status_var.set(f"Error: {e}")
            messagebox.showerror("Error", str(e))


def main():
    app = App()
    app.mainloop()


if __name__ == "__main__":
    main()
